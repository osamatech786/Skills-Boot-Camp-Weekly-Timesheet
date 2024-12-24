import streamlit as st
from datetime import datetime, date, timedelta
from docx import Document
from docx.shared import Inches
import numpy as np
import pandas as pd
from PIL import Image as PILImage
from io import BytesIO
from streamlit_drawable_canvas import st_canvas
import requests
import re
import os
from dotenv import load_dotenv
import msal
from openpyxl import load_workbook
import urllib.parse

# Set page configuration with a favicon
st.set_page_config(
    page_title="Skills Boot Camp Weekly Timesheet",
    page_icon="https://lirp.cdn-website.com/d8120025/dms3rep/multi/opt/social-image-88w.png", 
    layout="centered"  # "centered" or "wide"
)

# Load environment variables from .env file
load_dotenv()

# Fetch credentials from environment variables
CLIENT_ID = os.getenv("CLIENT_ID")
CLIENT_SECRET = os.getenv("CLIENT_SECRET")
TENANT_ID = os.getenv("TENANT_ID")
DRIVE_ID = os.getenv("DRIVE_ID")

# Authenticate and acquire an access token
def acquire_access_token():
    app = msal.ConfidentialClientApplication(
        client_id=CLIENT_ID,
        client_credential=CLIENT_SECRET,
        authority=f"https://login.microsoftonline.com/{TENANT_ID}",
    )
    result = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])
    if "access_token" in result:
        return result["access_token"]
    else:
        print("Failed to acquire token")
        print(result.get("error"))
        print(result.get("error_description"))
        exit()

ACCESS_TOKEN = acquire_access_token() 

# ========================
# Functions
# ========================
# add render support along with st.secret
def get_secret(key):
    try:
        load_dotenv()
        # Attempt to get the secret from environment variables
        secret = os.environ.get(key)
        if secret is None:
            raise ValueError("Secret not found in environment variables")
        return secret
    except (ValueError, TypeError) as e:
        # If an error occurs, fall back to Streamlit secrets
        if hasattr(st, 'secrets'):
            return st.secrets.get(key)
        # If still not found, return None or handle as needed
        return None
    
def upload_to_sharepoint(access_token, drive_id, parent_folder_path, attendance_sheet_file_path):
    headers = {"Authorization": f"Bearer {access_token}"}

    # URL encode the parent folder path
    encoded_parent_folder_path = urllib.parse.quote(parent_folder_path)

    # Fetch folders in the parent path
    parent_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{encoded_parent_folder_path}:/children"
    response = requests.get(parent_url, headers=headers)

    if response.status_code == 200:
        # Upload the attendance sheet using the multipart method
        upload_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{encoded_parent_folder_path}/{attendance_sheet_file_path}:/content"

        with open(attendance_sheet_file_path, "rb") as f:
            upload_response = requests.put(upload_url, headers=headers, data=f)
        return upload_response.status_code
    else:
        return f"Error fetching parent folder: {response.status_code} \nError details: {response.text}"

# Initialize session state for screen navigation
if 'page' not in st.session_state:
    st.session_state.page = 1
if 'learner_signature' not in st.session_state: 
    st.session_state.learner_signature = None
if 'declaration_date' not in st.session_state: 
    st.session_state.declaration_date = None    
if 'checkboxes' not in st.session_state: 
    st.session_state.checkboxes = []  # Store checkbox states for each attendance row
if 'attendance_checkboxes' not in st.session_state:
    st.session_state.attendance_checkboxes = []  # List of tuples [(am_present, am_absent, pm_present, pm_absent), ...]
if 'start_date' not in st.session_state:
    st.session_state.start_date = None
if 'end_date' not in st.session_state:
    st.session_state.end_date = None
    
# Function to load DOCX data, skipping header row in the second table
def load_docx_data():
    doc = Document(fr'resources/Skills Boot Camp Week {get_secret("week")} Group 1 Timesheet.docx')

    # Read the first paragraph for the weekly timesheet information
    weekly_timesheet_info = doc.paragraphs[1].text  # Get the first paragraph
    
    day, session_activity, facilitator, time, notes_comments = [], [], [], [], []
    attendance_data = []  # for second table
    
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            if table_idx == 0 and len(cells) == 5:
                day.append(cells[0])
                session_activity.append(cells[1])
                facilitator.append(cells[2])
                time.append(cells[3])
                notes_comments.append(cells[4])
            elif table_idx == 1 and len(cells) >= 4:
                if row_idx == 0:
                    continue  # Skip the header row
                while len(cells) < 5:
                    cells.append("")  # Add empty cell if columns are missing
                attendance_data.append(cells[:5])

    df1 = pd.DataFrame({
        "Day": day,
        "Session/Activity": session_activity,
        "Facilitator": facilitator,
        "Time": time,
        "Notes/Comments": notes_comments
    })
    df2 = pd.DataFrame(attendance_data, columns=["Day", "Date", "AM", "PM", "Learner Signature"])
    
    return weekly_timesheet_info, df1, df2

def is_signature_drawn(signature):
    if signature is None or not isinstance(signature, np.ndarray) or np.all(signature == 255):
        return False
    return True

def get_weekday_dates(start_date, end_date):
    """Generate start date, Tuesday, Wednesday, Thursday, and end date"""
    dates = {
        'start_date': start_date.strftime("%d/%m/%Y"),
        'tu_date': None,
        'we_date': None,
        'th_date': None,
        'end_date': end_date.strftime("%d/%m/%Y"),
    }

    current_date = start_date
    while current_date <= end_date:
        if current_date.weekday() == 1:  # Tuesday
            dates['tu_date'] = current_date.strftime("%d/%m/%Y")
        elif current_date.weekday() == 2:  # Wednesday
            dates['we_date'] = current_date.strftime("%d/%m/%Y")
        elif current_date.weekday() == 3:  # Thursday
            dates['th_date'] = current_date.strftime("%d/%m/%Y")
        current_date += timedelta(days=1)

    return dates

# Load data from DOCX
weekly_timesheet_info, df1, df2 = load_docx_data()

# First Screen: Display the Timesheet table
if st.session_state.page == 1:
    st.header(f'Skills Boot Camp Week {get_secret("week")} Timesheet')
    
    # Date input fields
    if isinstance(st.session_state.get("start_date"), str):
        st.session_state.start_date = datetime.strptime(st.session_state.get("start_date"), "%d/%m/%Y").date()
    # Date of Birth
    st.session_state.start_date = st.date_input(
        label="Start Date",  # Label for the field
        value=st.session_state.get("start_date"),  # Correctly access start_date from session state
        min_value=date(1900, 1, 1),  # Minimum selectable date
        max_value=date.today(),  # Maximum selectable date
        help="Choose a date",  # Tooltip text
        format='DD/MM/YYYY'
    )
    if st.session_state.start_date != None:
        st.session_state.start_date = st.session_state.start_date.strftime("%d/%m/%Y")

    if isinstance(st.session_state.get("end_date"), str):
        st.session_state.end_date = datetime.strptime(st.session_state.get("end_date"), "%d/%m/%Y").date()
    # Date of Birth
    st.session_state.end_date = st.date_input(
        label="End Date",  # Label for the field
        value=st.session_state.get("end_date"),  # Correctly access end_date from session state
        min_value=date(1900, 1, 1),  # Minimum selectable date
        max_value=date.today(),  # Maximum selectable date
        help="Choose a date",  # Tooltip text
        format='DD/MM/YYYY'
    )
    if st.session_state.end_date != None:
        st.session_state.end_date = st.session_state.end_date.strftime("%d/%m/%Y")
    
    # Check if both dates are provided
    if st.session_state.start_date and st.session_state.end_date:
        # Replace dates in weekly_timesheet_info
        weekly_timesheet_info = weekly_timesheet_info.replace("start_date", st.session_state.start_date)
        weekly_timesheet_info = weekly_timesheet_info.replace("end_date", st.session_state.end_date)
    else:
        st.warning("Please enter both start and end dates.")
        st.stop()  # Stop execution until both dates are entered


    st.text(weekly_timesheet_info)

    st.markdown(df1.to_html(index=False, header=False), unsafe_allow_html=True)
    
    if st.button("Next"):
        st.session_state.page = 2
        st.experimental_rerun()

# Second Screen: Learner Declaration and Attendance Table with Checkboxes
elif st.session_state.page == 2:    

    # Clear attendance checkboxes if returning to this page
    if not st.session_state.attendance_checkboxes:  
        st.session_state.attendance_checkboxes = [(False, False, False, False) for _ in range(len(df2))]

    st.subheader("Attendance Register Declaration (Monday - Friday)")
    st.session_state.learner_name = st.text_input("Enter your full name")
    if st.session_state.learner_name:
        if not st.session_state.learner_name.replace(" ", "").isalpha():
            st.warning("Please enter a valid name (letters and spaces only)")
        elif len(st.session_state.learner_name.strip()) < 2:
            st.warning("Name must be at least 2 characters long")

    st.markdown(f"I, {st.session_state.learner_name} confirm I have attended the scheduled sessions from **{st.session_state.start_date}** to **{st.session_state.end_date}** "
                "as outlined in the weekly timetable. I understand that accurate attendance is important for the completion of this programme.")
    
    # Custom header without checkboxes
    st.markdown("### Attendance Table")
    header_cols = st.columns([1, 1, 1, 1, 1])
    header_cols[0].write("Day")
    header_cols[1].write("Date")
    header_cols[2].write("AM")
    header_cols[3].write("PM")
    header_cols[4].write("Learner Signature")


    # Track checkbox states in session state to retain selections
    checkboxes = []
    
    weekday_dates = get_weekday_dates(datetime.strptime(st.session_state.get("start_date"), "%d/%m/%Y").date(), datetime.strptime(st.session_state.get("end_date"), "%d/%m/%Y").date())
    for idx, row in df2.iterrows():
        row_cols = st.columns([1, 1, 1, 1, 1])
        row_cols[0].write(row["Day"])
        # row_cols[1].write(row["Date"])

        # Use the weekday_dates dictionary to get the date value
        # Assuming you want to map 'Date' in df2 to the keys from weekday_dates
        if row["Date"] == "start_date":
            row_cols[1].write(weekday_dates['start_date'])
        elif row["Date"] == "tu_date":
            row_cols[1].write(weekday_dates['tu_date'])
        elif row["Date"] == "we_date":
            row_cols[1].write(weekday_dates['we_date'])
        elif row["Date"] == "th_date":
            row_cols[1].write(weekday_dates['th_date'])
        elif row["Date"] == "end_date":
            row_cols[1].write(weekday_dates['end_date'])
        
        # Checkboxes for AM
        am_present = row_cols[2].checkbox("Present (AM)", key=f"am_present_{idx}", value=st.session_state.attendance_checkboxes[idx][0])
        am_absent = row_cols[2].checkbox("Absent (AM)", key=f"am_absent_{idx}", value=st.session_state.attendance_checkboxes[idx][1])

        # Checkboxes for PM
        pm_present = row_cols[3].checkbox("Present (PM)", key=f"pm_present_{idx}", value=st.session_state.attendance_checkboxes[idx][2])
        pm_absent = row_cols[3].checkbox("Absent (PM)", key=f"pm_absent_{idx}", value=st.session_state.attendance_checkboxes[idx][3])

        # Store checkbox states in a tuple
        st.session_state.attendance_checkboxes[idx] = (am_present, am_absent, pm_present, pm_absent)

        # Learner Signature Checkbox
        checked = row_cols[4].checkbox("Signature", key=f"signature_{idx}")
        checkboxes.append(checked)
    
    # Update checkboxes in session state
    st.session_state.checkboxes = checkboxes

    # Signature Section
    st.subheader("Learner Declaration")
    st.write("I confirm that the information above is correct and that my attendance has been accurately recorded for this week.")
    st.text("Please draw your signature below:")

    # Signature Box
    canvas_result = st_canvas(
        fill_color="rgba(255, 255, 255, 1)",  
        stroke_width=5,
        stroke_color="rgb(0, 0, 0)",
        background_color="white",
        width=400,
        height=150,
        drawing_mode="freedraw",
        key="canvas",
    )
    st.session_state.learner_signature = canvas_result.image_data

    declaration_date = date.today().strftime("%d-%m-%Y")
    st.write(f"Date: **{declaration_date}**")    

    if st.button("Submit"):
        
        valid_attendance = True
        for idx, (am_present, am_absent, pm_present, pm_absent) in enumerate(st.session_state.attendance_checkboxes):
            print(f"Day {idx + 1}: AM Present: {am_present}, AM Absent: {am_absent}, PM Present: {pm_present}, PM Absent: {pm_absent}")
            if not (am_present or pm_present or am_absent or pm_absent):  # At least one AM or PM must be checked
                valid_attendance = False
                print(f"Invalid attendance on Day {idx + 1}")
                break

        print(f"Valid attendance: {valid_attendance}")

        if valid_attendance:
            if is_signature_drawn(st.session_state.learner_signature) and st.session_state.learner_name:
                filled_doc = Document(fr'resources/Skills Boot Camp Week {get_secret("week")} Group 1 Timesheet.docx')
                
                for paragraph in filled_doc.paragraphs:
                    if 'start_date' in paragraph.text:
                        paragraph.text = paragraph.text.replace('start_date', str(st.session_state.start_date))
                    if 'end_date' in paragraph.text:
                        paragraph.text = paragraph.text.replace('end_date', str(st.session_state.end_date)                        )
                    if 'learner_name' in paragraph.text:
                        paragraph.text = paragraph.text.replace('learner_name', str(st.session_state.learner_name))
                    if 'date' in paragraph.text:
                        paragraph.text = paragraph.text.replace('date', declaration_date)
                    # Replace the placeholder text with the signature image
                    if 'learner_signature' in paragraph.text:
                        paragraph.text = paragraph.text.replace('learner_signature', "")
                        
                        # Create a new run within the paragraph to add the image
                        run = paragraph.add_run()
                        signature_image = PILImage.fromarray(st.session_state.learner_signature.astype('uint8'), 'RGBA')
                        signature_image_path = "learner_signature.png"
                        signature_image.save(signature_image_path)
                        
                        # Insert the image into the document in place of the placeholder
                        run.add_picture(signature_image_path, width=Inches(2))  # Adjust width as needed

                # Insert signature image in the "Learner Signature" cell for each checked row
                for table_idx, table in enumerate(filled_doc.tables):
                    if table_idx == 1:  # Ensure we're modifying the attendance table
                        for row_idx, row in enumerate(table.rows[1:]):  # Skip the header row

                            # Replace keys with date values
                            day_text = row.cells[1].text  # Assuming the day is in the first cell
                            if day_text == "start_date":
                                row.cells[1].text = weekday_dates['start_date']
                            elif day_text == "tu_date":
                                row.cells[1].text = weekday_dates['tu_date']
                            elif day_text == "we_date":
                                row.cells[1].text = weekday_dates['we_date']
                            elif day_text == "th_date":
                                row.cells[1].text = weekday_dates['th_date']
                            elif day_text == "end_date":
                                row.cells[1].text = weekday_dates['end_date']

                            # Get the checkboxes state for the current row
                            am_present_checked = st.session_state.attendance_checkboxes[row_idx][0]  # AM Present
                            am_absent_checked = st.session_state.attendance_checkboxes[row_idx][1]   # AM Absent
                            pm_present_checked = st.session_state.attendance_checkboxes[row_idx][2]  # PM Present
                            pm_absent_checked = st.session_state.attendance_checkboxes[row_idx][3]   # PM Absent
                            
                            # Handle AM placeholders
                            am_text = row.cells[2].text  # AM is in cell 2
                            if am_present_checked:
                                am_text = am_text.replace('[am_pr]', '✔').replace('[am_ab]', ' ')
                            else:
                                am_text = am_text.replace('[am_pr]', ' ').replace('[am_ab]', '✔' if am_absent_checked else ' ')

                            row.cells[2].text = am_text
                            
                            # Handle PM placeholders
                            pm_text = row.cells[3].text  # PM is in cell 3
                            if pm_present_checked:
                                pm_text = pm_text.replace('[pm_pr]', '✔').replace('[pm_ab]', ' ')
                            else:
                                pm_text = pm_text.replace('[pm_pr]', ' ').replace('[pm_ab]', '✔' if pm_absent_checked else ' ')

                            row.cells[3].text = pm_text

                            # "Learner Signature" column
                            cell = row.cells[4]  
                            if st.session_state.checkboxes[row_idx]:  # If checkbox was checked
                                # Clear text and add image
                                cell.text = ""
                                run = cell.paragraphs[0].add_run()
                                run.add_picture(signature_image_path, width=Inches(0.5))  # Adjust width as needed
                            else:
                                cell.text = "Absent"  # Clear any placeholder text if unchecked

                # After using the signature image
                try:
                    if os.path.exists(signature_image_path):
                        os.remove(signature_image_path)
                except Exception as e:
                    st.warning(f"Could not remove temporary signature file: {e}")
                
                # Generate a unique file name based on the learner's name
                safe_learner_name = re.sub(r'\W+', '_', st.session_state.learner_name)
                filled_doc_path = f'Timesheet_w{get_secret("week")}_{safe_learner_name}.docx'
                # Add error handling for document saving
                try:
                    filled_doc.save(filled_doc_path)
                except Exception as e:
                    st.error(f"Error saving document: {e}")

                # with open(filled_doc_path, "rb") as file:
                #     st.download_button(f"Download Filled Timesheet for {st.session_state.learner_name}", file, filled_doc_path)
                
                with st.spinner('Submitting your timesheet...'):
                    # Upload to share point
                    parent_folder_path = get_secret("PARENT_FOLDER_PATH")
                    status_code=upload_to_sharepoint(ACCESS_TOKEN, DRIVE_ID, parent_folder_path, filled_doc_path)
                    if status_code == 200:
                        st.warning(f"Timesheet already exist with the same name.")
                    elif status_code == 201:
                        st.success(f"Timesheet submitted successfully!")
                    elif status_code == 400:
                        st.error(f"Error submitting timesheet!")
                    else:
                        st.error(status_code)
            else:
                st.warning("Please enter your name & draw the signature!")
        else:
            st.warning("Please ensure at least one attendance checkbox (AM or PM) is checked for each day!")

    if st.button("Back"):
        st.session_state.page = 1
        st.experimental_rerun()


# python -m streamlit run app.py --server.port 8505
# Dev : https://linkedin.com/in/osamatech786